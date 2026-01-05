from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from django.http import HttpResponse

def export_pdf(summaries):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="summary.pdf"'
    c = canvas.Canvas(response, pagesize=A4)

    y = 800
    for item in summaries:
        c.drawString(40, y, f"Page {item['page']}")
        y -= 20
        c.drawString(40, y, item["summary"][:1000])
        y -= 40
        if y < 100:
            c.showPage()
            y = 800

    c.save()
    return response

def export_docx(summaries):
    doc = Document()
    doc.add_heading("PDF Summary", level=1)

    for item in summaries:
        doc.add_heading(f"Page {item['page']}", level=2)
        doc.add_paragraph(item["summary"])

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="summary.docx"'
    doc.save(response)
    return response
